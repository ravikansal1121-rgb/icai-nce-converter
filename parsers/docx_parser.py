"""
Parse T-format financial statements from Word (.docx) files.
Extracts tables and falls back to paragraph text parsing.
Returns the same structure as xlsx_parser: list of dicts.
"""
import re

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None


def _clean_amount(val):
    """Convert a string/number value to a float amount."""
    if val is None:
        return 0.0
    if isinstance(val, (int, float)):
        return float(val)
    s = str(val).strip()
    if not s or s in ('-', '—', '–', 'nil', 'Nil', 'NIL', '0'):
        return 0.0
    neg = False
    if s.startswith('(') and s.endswith(')'):
        neg = True
        s = s[1:-1]
    s = s.replace(',', '').replace('₹', '').replace('Rs.', '').replace('Rs', '').strip()
    try:
        v = float(s)
        return -v if neg else v
    except ValueError:
        return 0.0


def _clean_name(val):
    """Clean an account name string."""
    if val is None:
        return ''
    s = str(val).strip()
    s = re.sub(r'^(To|By)\s+', '', s, flags=re.IGNORECASE).strip()
    s = re.sub(r'\s+', ' ', s)
    return s


def _is_skip_name(name):
    """Check if a name should be skipped (totals, headers, etc.)."""
    skip_patterns = [
        r'^total', r'^grand total', r'^net profit', r'^net loss',
        r'^gross profit', r'^gross loss', r'^balance c/d',
        r'^balance b/d', r'^as per', r'^amount',
        r'^particulars', r'^dr\.?$', r'^cr\.?$',
        r'^assets?$', r'^liabilit', r'^equity',
        r'^current assets?$', r'^non.?current', r'^fixed assets?$',
    ]
    lower = name.lower()
    return any(re.match(p, lower) for p in skip_patterns)


def _extract_year_from_text(text):
    """Try to find a financial year pattern in text."""
    m = re.search(r'(\d{4})\s*[-–—]\s*(\d{2,4})', text)
    if m:
        y1 = m.group(1)
        y2 = m.group(2)
        if len(y2) == 2:
            y2 = y1[:2] + y2
        return f'{y1}-{y2[-2:]}'
    m = re.search(r'31\s*(st)?\s*march\s*,?\s*(\d{4})', text, re.IGNORECASE)
    if m:
        y = int(m.group(2))
        return f'{y-1}-{str(y)[-2:]}'
    return ''


def _detect_section(text):
    """Detect which section a header row belongs to."""
    t = text.lower().strip()
    trading_kw = ['trading', 'trading account', 'trading and profit', 'trading & profit',
                  'profit and loss', 'profit & loss', 'income and expenditure']
    bs_kw = ['balance sheet', 'balancesheet']
    capital_kw = ['capital account', 'capital a/c']

    for kw in trading_kw:
        if kw in t:
            return 'trading_pl'
    for kw in bs_kw:
        if kw in t:
            return 'balance_sheet'
    for kw in capital_kw:
        if kw in t:
            return 'capital'
    return None


def _parse_t_format_table(table):
    """
    Parse a T-format table (side-by-side Dr/Cr or Liabilities/Assets).
    Returns (left_items, right_items) where each is a list of {'name': ..., 'amount': ...}.
    """
    rows = table.rows
    if len(rows) < 2:
        return [], []

    num_cols = len(rows[0].cells)
    if num_cols < 2:
        return [], []

    # Find the midpoint column
    mid = num_cols // 2

    left_items = []
    right_items = []

    for row in rows[1:]:  # skip header row
        cells = row.cells
        # Left side
        left_name = ''
        left_amt = 0.0
        for c in range(0, mid):
            text = cells[c].text.strip() if c < len(cells) else ''
            if text:
                amt = _clean_amount(text)
                if amt != 0 and not left_name:
                    # Amount without a name, skip
                    continue
                if amt != 0:
                    left_amt = amt
                else:
                    cleaned = _clean_name(text)
                    if cleaned and not left_name:
                        left_name = cleaned

        if left_name and not _is_skip_name(left_name):
            # Try to find amount in remaining left cells
            if left_amt == 0:
                for c in range(1, mid):
                    if c < len(cells):
                        left_amt = _clean_amount(cells[c].text.strip())
                        if left_amt != 0:
                            break
            if left_amt != 0:
                left_items.append({'name': left_name, 'amount': abs(left_amt)})

        # Right side
        right_name = ''
        right_amt = 0.0
        for c in range(mid, num_cols):
            text = cells[c].text.strip() if c < len(cells) else ''
            if text:
                amt = _clean_amount(text)
                if amt != 0 and not right_name:
                    continue
                if amt != 0:
                    right_amt = amt
                else:
                    cleaned = _clean_name(text)
                    if cleaned and not right_name:
                        right_name = cleaned

        if right_name and not _is_skip_name(right_name):
            if right_amt == 0:
                for c in range(mid + 1, num_cols):
                    if c < len(cells):
                        right_amt = _clean_amount(cells[c].text.strip())
                        if right_amt != 0:
                            break
            if right_amt != 0:
                right_items.append({'name': right_name, 'amount': abs(right_amt)})

    return left_items, right_items


def _parse_capital_from_table(table):
    """Extract capital account details from a table."""
    cap = {
        'opening': 0, 'capital_introduced': 0, 'net_profit': 0,
        'interest_on_capital': 0, 'drawings': 0, 'closing': 0,
    }
    for row in table.rows:
        name = ''
        amt = 0.0
        for cell in row.cells:
            text = cell.text.strip()
            if not text:
                continue
            a = _clean_amount(text)
            if a != 0:
                amt = a
            elif text:
                name = text.lower()

        if not name:
            continue
        if 'opening' in name or 'balance b/d' in name:
            cap['opening'] = abs(amt)
        elif 'capital introduced' in name or 'additional capital' in name:
            cap['capital_introduced'] = abs(amt)
        elif 'net profit' in name or 'profit for' in name:
            cap['net_profit'] = abs(amt)
        elif 'interest on capital' in name:
            cap['interest_on_capital'] = abs(amt)
        elif 'drawing' in name or 'withdrawal' in name:
            cap['drawings'] = abs(amt)
        elif 'closing' in name or 'balance c/d' in name:
            cap['closing'] = abs(amt)

    if cap['closing'] == 0 and cap['opening'] > 0:
        cap['closing'] = (cap['opening'] + cap['capital_introduced'] +
                          cap['net_profit'] + cap['interest_on_capital'] - cap['drawings'])
    return cap


def _infer_constitution(data):
    """Infer whether entity is proprietorship or partnership."""
    all_names = []
    for section in ['trading_pl', 'balance_sheet']:
        if section in data and isinstance(data[section], dict):
            for side_key in ['debit', 'credit', 'liabilities', 'assets']:
                if side_key in data[section]:
                    all_names.extend([i['name'].lower() for i in data[section][side_key]])

    entity = data.get('entity_name', '').lower()
    if '& co' in entity or 'and co' in entity:
        return 'partnership'
    for n in all_names:
        if 'partner' in n:
            return 'partnership'
    return 'proprietorship'


def _extract_entity_from_paragraphs(paragraphs):
    """Try to find entity name from the first few paragraphs."""
    skip = ['trading', 'profit', 'balance', 'capital', 'for the',
            'as at', 'year ended', 'dr', 'cr', 'particulars', 'amount']
    for para in paragraphs[:10]:
        text = para.text.strip()
        if not text or len(text) < 4:
            continue
        t_lower = text.lower()
        if any(s in t_lower for s in skip):
            continue
        if re.match(r'^[\d,.\s]+$', text):
            continue
        return text
    return 'Entity Name'


def parse_docx(filepath):
    """
    Parse a Word document containing T-format financial statements.
    Returns a list of parsed data dicts (same structure as xlsx_parser).
    """
    if DocxDocument is None:
        return []

    # For .doc files, try to open as .docx
    ext = filepath.rsplit('.', 1)[-1].lower() if '.' in filepath else ''
    if ext == 'doc':
        try:
            doc = DocxDocument(filepath)
        except Exception:
            return []
    else:
        try:
            doc = DocxDocument(filepath)
        except Exception:
            return []

    # Try to extract entity name and year from paragraphs
    entity_name = _extract_entity_from_paragraphs(doc.paragraphs)
    year = ''
    for para in doc.paragraphs[:15]:
        y = _extract_year_from_text(para.text)
        if y:
            year = y
            break

    # If no year found in paragraphs, check table cells
    if not year:
        for table in doc.tables[:3]:
            for row in table.rows[:5]:
                for cell in row.cells:
                    y = _extract_year_from_text(cell.text)
                    if y:
                        year = y
                        break
                if year:
                    break
            if year:
                break

    data = {
        'entity_name': entity_name,
        'fy': year,
        'source_sheet': 'docx',
    }

    # Process tables
    for table in doc.tables:
        # Check first row/cells for section header
        header_text = ''
        if table.rows:
            for cell in table.rows[0].cells:
                header_text += ' ' + cell.text
        # Also check the paragraph immediately before this table
        section = _detect_section(header_text)

        if section == 'trading_pl':
            debit, credit = _parse_t_format_table(table)
            if debit or credit:
                data['trading_pl'] = {'debit': debit, 'credit': credit}
        elif section == 'balance_sheet':
            liab, assets = _parse_t_format_table(table)
            if liab or assets:
                data['balance_sheet'] = {'liabilities': liab, 'assets': assets}
        elif section == 'capital':
            data['capital_account'] = _parse_capital_from_table(table)
        else:
            # Try to detect from content if no header
            if 'trading_pl' not in data:
                debit, credit = _parse_t_format_table(table)
                if debit or credit:
                    # Check if items look like trading/PL items
                    all_names = [i['name'].lower() for i in debit + credit]
                    has_trading = any(kw in ' '.join(all_names) for kw in
                                     ['sales', 'purchase', 'salary', 'rent', 'stock', 'depreciation'])
                    if has_trading:
                        data['trading_pl'] = {'debit': debit, 'credit': credit}
                        continue

            if 'balance_sheet' not in data:
                left, right = _parse_t_format_table(table)
                if left or right:
                    all_names = [i['name'].lower() for i in left + right]
                    has_bs = any(kw in ' '.join(all_names) for kw in
                                ['capital', 'bank', 'cash', 'debtor', 'creditor', 'loan',
                                 'fixed deposit', 'sundry'])
                    if has_bs:
                        data['balance_sheet'] = {'liabilities': left, 'assets': right}

    # If no tables found, try paragraph-based extraction
    if not any(k in data for k in ['trading_pl', 'balance_sheet', 'capital_account']):
        items = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            # Try to find "Name ... Amount" patterns
            parts = re.split(r'\s{2,}|\t+', text)
            if len(parts) >= 2:
                name = _clean_name(parts[0])
                amt = _clean_amount(parts[-1])
                if name and amt != 0 and not _is_skip_name(name):
                    items.append({'name': name, 'amount': abs(amt)})
        if items:
            data['raw_items'] = items

    # Only return if we found financial data
    if any(k in data for k in ['trading_pl', 'balance_sheet', 'capital_account', 'raw_items']):
        data['constitution'] = _infer_constitution(data)
        data['proprietor_name'] = entity_name if data['constitution'] == 'proprietorship' else ''
        return [data]

    return []
